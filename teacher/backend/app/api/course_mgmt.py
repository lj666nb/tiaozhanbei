"""
课程管理 API — 全动态后端驱动，无前端硬编码

提供：枚举配置 / 课程CRUD / 教师管理 / 学期管理
"""

from __future__ import annotations

import io
import json
import uuid
from datetime import datetime
from pathlib import Path
from urllib.parse import quote

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

from app.models.schemas import APIResponse

router = APIRouter(prefix="/api/course-mgmt", tags=["课程管理"])

# ── 数据文件存储（生产环境应使用 DB） ─────────────────
# 统一使用 /app/data/（Docker 持久化卷），本地开发时自动创建 data/ 目录
import os as _os
_DATA_ROOT = _os.environ.get("DATA_DIR", str(Path(__file__).parent.parent.parent / "data"))
DATA_DIR = Path(_DATA_ROOT)
COURSES_FILE = DATA_DIR / "courses.json"
ENUMS_FILE = DATA_DIR / "course_enums.json"
TEACHERS_FILE = DATA_DIR / "course_teachers.json"

DATA_DIR.mkdir(parents=True, exist_ok=True)


def _read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _write_json(path: Path, data: dict) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


# ═══════════════════════════════════════════════════════
# 枚举配置
# ═══════════════════════════════════════════════════════

@router.get("/enums", response_model=APIResponse)
async def get_enums():
    """获取全部动态枚举配置（所有下拉选项从此接口读取）。"""
    enums = _read_json(ENUMS_FILE)
    defaults = {
        "semesters": [
            {"value": "2026春季", "label": "2026 春季"},
            {"value": "2025秋季", "label": "2025 秋季"},
            {"value": "2025春季", "label": "2025 春季"},
        ],
        "statuses": [
            {"value": "进行中", "label": "进行中", "color": "#1677ff", "type": "processing"},
            {"value": "已过半", "label": "已过半", "color": "#fa8c16", "type": "warning"},
            {"value": "已结课", "label": "已结课", "color": "#52c41a", "type": "success"},
        ],
        "course_categories": [
            {"value": "专业核心", "label": "专业核心"},
            {"value": "专业选修", "label": "专业选修"},
            {"value": "通识基础", "label": "通识基础"},
        ],
        "progress_thresholds": {
            "half": 50,      # 进度 ≥50% 标记为"已过半"
            "complete": 100, # 进度=100% 标记为"已结课"
            "max_hours": 64, # 课时上限
        },
    }
    # 合并默认值与已保存的自定义配置
    for k, v in defaults.items():
        if k not in enums or not enums[k]:
            enums[k] = v
    _write_json(ENUMS_FILE, enums)
    return APIResponse(success=True, data=enums)


@router.post("/enums/semesters", response_model=APIResponse)
async def add_semester(data: dict):
    """新增学期（如"2026秋季"）。"""
    enums = _read_json(ENUMS_FILE)
    semesters = enums.get("semesters", [])
    value = data.get("value", "").strip()
    label = data.get("label", value)
    if not value:
        raise HTTPException(status_code=400, detail="学期值不能为空")
    if any(s["value"] == value for s in semesters):
        raise HTTPException(status_code=400, detail="学期已存在")
    semesters.append({"value": value, "label": label})
    enums["semesters"] = semesters
    _write_json(ENUMS_FILE, enums)
    return APIResponse(success=True, message=f"已添加学期「{label}」", data=enums)


# ═══════════════════════════════════════════════════════
# 教师管理
# ═══════════════════════════════════════════════════════

@router.get("/teachers", response_model=APIResponse)
async def list_teachers():
    """获取教师列表。"""
    teachers = _read_json(TEACHERS_FILE)
    lst = teachers.get("list", [])
    # 首次使用自动从课程数据中提取教师
    if not lst:
        courses = _read_json(COURSES_FILE)
        course_list = courses.get("list", [])
        seen = set()
        for c in course_list:
            t = c.get("teacher", "").strip()
            if t and t not in seen:
                seen.add(t)
                lst.append({"name": t, "title": "", "id": str(uuid.uuid4())[:8]})
        if lst:
            teachers["list"] = lst
            _write_json(TEACHERS_FILE, teachers)
    return APIResponse(success=True, data={"teachers": lst, "total": len(lst)})


@router.post("/teachers", response_model=APIResponse)
async def add_teacher(data: dict):
    """新增教师。"""
    name = data.get("name", "").strip()
    title = data.get("title", "讲师").strip()
    if not name:
        raise HTTPException(status_code=400, detail="教师姓名不能为空")
    teachers = _read_json(TEACHERS_FILE)
    lst = teachers.get("list", [])
    if any(t["name"] == name for t in lst):
        raise HTTPException(status_code=400, detail="教师已存在")
    lst.append({"name": name, "title": title, "id": str(uuid.uuid4())[:8]})
    teachers["list"] = lst
    _write_json(TEACHERS_FILE, teachers)
    return APIResponse(success=True, message=f"已添加教师「{name}」")


@router.delete("/teachers/{teacher_id}", response_model=APIResponse)
async def delete_teacher(teacher_id: str):
    """删除教师。"""
    teachers = _read_json(TEACHERS_FILE)
    before = len(teachers.get("list", []))
    teachers["list"] = [t for t in teachers["list"] if t["id"] != teacher_id]
    if len(teachers["list"]) == before:
        raise HTTPException(status_code=404, detail="教师不存在")
    _write_json(TEACHERS_FILE, teachers)
    return APIResponse(success=True, message="已删除")


# ═══════════════════════════════════════════════════════
# 课程 CRUD
# ═══════════════════════════════════════════════════════

class CourseCreate(BaseModel):
    name: str
    code: str = ""
    teacher: str = ""
    semester: str = "2026春季"
    category: str = "专业核心"
    description: str = ""
    max_hours: int = 48
    status: str = "进行中"


class CourseUpdate(BaseModel):
    name: str
    code: str = ""
    teacher: str = ""
    semester: str = "2026春季"
    category: str = "专业核心"
    description: str = ""
    max_hours: int = 48
    status: str = ""
    sessions: int | None = None
    updated_by: str = ""


# ── 种子课程模板（每次加载时自动恢复，锁定不可修改） ──
def _seed_courses() -> list[dict]:
    now = datetime.now().isoformat()[:19]
    return [
        {"id": "c001", "name": "机器学习", "code": "CS401", "semester": "2026春季", "category": "专业核心",
         "teacher": "张教授", "max_hours": 48, "progress": 65, "status": "已过半",
         "student_list": [{"id": "s1", "name": "张三", "student_id": "2024001"}, {"id": "s2", "name": "李四", "student_id": "2024002"}, {"id": "s3", "name": "王五", "student_id": "2024003"}],
         "sessions": 31, "_source": "seed", "created_at": now, "updated_at": now},
        {"id": "c002", "name": "深度学习", "code": "CS402", "semester": "2026春季", "category": "专业核心",
         "teacher": "李教授", "max_hours": 64, "progress": 42, "status": "进行中",
         "student_list": [{"id": "s4", "name": "赵六", "student_id": "2024004"}, {"id": "s5", "name": "孙七", "student_id": "2024005"}],
         "sessions": 27, "_source": "seed", "created_at": now, "updated_at": now},
        {"id": "c003", "name": "自然语言处理", "code": "CS403", "semester": "2026春季", "category": "专业选修",
         "teacher": "王教授", "max_hours": 32, "progress": 88, "status": "已过半",
         "student_list": [{"id": "s1", "name": "张三", "student_id": "2024001"}, {"id": "s6", "name": "周八", "student_id": "2024006"}, {"id": "s7", "name": "吴九", "student_id": "2024007"}],
         "sessions": 28, "_source": "seed", "created_at": now, "updated_at": now},
        {"id": "c004", "name": "计算机视觉", "code": "CS404", "semester": "2025秋季", "category": "专业核心",
         "teacher": "赵教授", "max_hours": 48, "progress": 100, "status": "已结课",
         "student_list": [{"id": "s2", "name": "李四", "student_id": "2024002"}, {"id": "s4", "name": "赵六", "student_id": "2024004"}],
         "sessions": 48, "_source": "seed", "created_at": now, "updated_at": now},
        {"id": "c005", "name": "数据挖掘", "code": "CS405", "semester": "2025秋季", "category": "专业选修",
         "teacher": "张教授", "max_hours": 32, "progress": 100, "status": "已结课",
         "student_list": [{"id": "s3", "name": "王五", "student_id": "2024003"}, {"id": "s5", "name": "孙七", "student_id": "2024005"}, {"id": "s8", "name": "郑十", "student_id": "2024008"}],
         "sessions": 32, "_source": "seed", "created_at": now, "updated_at": now},
    ]

_SEED_TEACHERS = [
    {"name": "张教授", "title": "教授", "id": "seed_t1"},
    {"name": "李教授", "title": "教授", "id": "seed_t2"},
    {"name": "王教授", "title": "副教授", "id": "seed_t3"},
    {"name": "赵教授", "title": "教授", "id": "seed_t4"},
]


@router.get("/courses", response_model=APIResponse)
async def list_courses(semester: str = ""):
    """获取课程列表（种子数据仅首次初始化，之后由用户管理）。"""
    courses = _read_json(COURSES_FILE)
    lst = courses.get("list", [])

    # ── 一次性种子初始化 ──
    # 仅当 courses.json 中无 _SEED_DONE_ 标记时才写入种子数据
    if not courses.get("_SEED_DONE_"):
        existing_ids = {c["id"] for c in lst}
        for seed_course in _seed_courses():
            if seed_course["id"] not in existing_ids:
                lst.append(seed_course)
        courses["list"] = lst
        courses["_SEED_DONE_"] = True
        _write_json(COURSES_FILE, courses)

    # 同步种子教师（仅首次）
    tdata = _read_json(TEACHERS_FILE)
    if not tdata.get("_SEED_DONE_"):
        existing = {t["id"] for t in tdata.get("list", [])}
        for st in _SEED_TEACHERS:
            if st["id"] not in existing:
                tdata.setdefault("list", []).append(st)
        tdata["_SEED_DONE_"] = True
        _write_json(TEACHERS_FILE, tdata)

    if semester:
        lst = [c for c in lst if c.get("semester") == semester]
    # 动态重算进度（基于 session_details 实际课时）
    for c in lst:
        details = c.get("session_details", [])
        if details:
            total_h = sum(d.get("hours", 0) for d in details)
            max_h = max(c.get("max_hours", 48), 1)
            c["sessions"] = total_h
            c["progress"] = min(round(total_h / max_h * 100), 100)
            c["status"] = "已结课" if c["progress"] >= 100 else ("已过半" if c["progress"] >= 50 else "进行中")
        else:
            stored_sessions = c.get("sessions", 0)
            stored_progress = c.get("progress", 0)
            max_h = max(c.get("max_hours", 48), 1)
            if stored_sessions > 0:
                c["progress"] = min(round(stored_sessions / max_h * 100), 100)
            elif stored_progress > 0:
                c["sessions"] = min(round(stored_progress / 100 * max_h), max_h)
            c["status"] = "已结课" if c["progress"] >= 100 else ("已过半" if c["progress"] >= 50 else "进行中")
        c["students"] = len(c.get("student_list", []))
    lst.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return APIResponse(success=True, data={"total": len(lst), "items": lst})


@router.get("/courses/{course_id}", response_model=APIResponse)
async def get_course(course_id: str):
    """获取单门课程详情。"""
    courses = _read_json(COURSES_FILE)
    for c in courses.get("list", []):
        if c["id"] == course_id:
            return APIResponse(success=True, data=c)
    raise HTTPException(status_code=404, detail="课程不存在")


@router.post("/courses", response_model=APIResponse)
async def create_course(data: CourseCreate):
    """新增课程（含编号重复校验）。"""
    courses = _read_json(COURSES_FILE)
    lst = courses.get("list", [])

    # 编号重复校验
    new_code = data.code.strip()
    if new_code:
        for c in lst:
            if c.get("code") == new_code:
                raise HTTPException(status_code=409, detail=f"课程编号「{new_code}」已被「{c['name']}」使用")

    course_id = str(uuid.uuid4())[:8]
    now = datetime.now().isoformat()[:19]
    status = data.status.strip() or "进行中"

    max_h = max(1, min(data.max_hours, 64))
    course = {
        "id": course_id,
        "name": data.name.strip(),
        "code": new_code or f"AI{course_id}",
        "class_code": "",  # 班级编号（学生用此加入）
        "teacher": data.teacher.strip(),
        "semester": data.semester,
        "category": data.category,
        "description": data.description.strip(),
        "max_hours": max_h,
        "students": 0,
        "sessions": 0,
        "progress": 0,
        "status": status,
        "color": "#1677ff",
        "student_list": [],
        "_source": "user",
        "created_at": now,
        "updated_at": now,
        "updated_by": "",
    }
    lst.append(course)
    courses["list"] = lst
    _write_json(COURSES_FILE, courses)
    return APIResponse(success=True, message=f"课程「{course['name']}」创建成功", data=course)


@router.put("/courses/{course_id}", response_model=APIResponse)
async def update_course(course_id: str, data: CourseUpdate):
    """更新课程（编号重复校验 / 状态联动进度 / 编辑历史）。"""
    courses = _read_json(COURSES_FILE)
    lst = courses.get("list", [])

    # 课程编号重复校验
    new_code = data.code.strip()
    if new_code:
        for c in lst:
            if c["id"] != course_id and c.get("code") == new_code:
                raise HTTPException(status_code=409, detail=f"课程编号「{new_code}」已被「{c['name']}」使用")

    for i, c in enumerate(lst):
        if c["id"] == course_id:
            old_max_hours = c.get("max_hours", 48)
            lst[i]["name"] = data.name.strip()
            lst[i]["code"] = new_code or lst[i]["code"]
            lst[i]["teacher"] = data.teacher.strip()
            lst[i]["semester"] = data.semester
            lst[i]["category"] = data.category
            lst[i]["description"] = data.description.strip()
            lst[i]["max_hours"] = max(1, min(data.max_hours, 64))
            lst[i]["updated_at"] = datetime.now().isoformat()[:19]
            if data.updated_by.strip():
                lst[i]["updated_by"] = data.updated_by.strip()

            # 显式传入 sessions 时更新已授课时
            if data.sessions is not None:
                lst[i]["sessions"] = max(0, data.sessions)

            # 状态联动：根据 max_hours 变更和显式状态决定是否重算
            max_hours_changed = old_max_hours != lst[i]["max_hours"]
            new_status = data.status.strip()

            # max_hours 变更时基于 session_details / sessions 重算进度
            if max_hours_changed:
                _recalc_course_progress(lst[i])

            if new_status:
                lst[i]["status"] = new_status
                if new_status == "已结课":
                    lst[i]["progress"] = 100
                    lst[i]["sessions"] = lst[i].get("max_hours", 48)
                elif new_status == "进行中" and lst[i].get("progress", 0) >= 100:
                    lst[i]["progress"] = min(lst[i].get("progress", 99), 99)
            else:
                # 未显式设状态 → 基于 session_details 重算进度和状态
                _recalc_course_progress(lst[i])

            # 同步学生数
            lst[i]["students"] = len(lst[i].get("student_list", []))

            courses["list"] = lst
            _write_json(COURSES_FILE, courses)
            return APIResponse(success=True, message="已更新", data=lst[i])
    raise HTTPException(status_code=404, detail="课程不存在")


@router.delete("/courses/{course_id}", response_model=APIResponse)
async def delete_course(course_id: str):
    """删除课程。"""
    courses = _read_json(COURSES_FILE)
    lst = courses.get("list", [])
    if not lst:
        raise HTTPException(status_code=404, detail="课程不存在")
    before = len(lst)
    lst = [c for c in lst if c["id"] != course_id]
    if len(lst) == before:
        raise HTTPException(status_code=404, detail="课程不存在")
    courses["list"] = lst
    _write_json(COURSES_FILE, courses)
    return APIResponse(success=True, message="已删除")


@router.post("/courses/{course_id}/add-session", response_model=APIResponse)
async def add_session(course_id: str):
    """新增一节课时（默认1小时），自动基于 session_details 重算进度和状态。"""
    courses = _read_json(COURSES_FILE)
    lst = courses.get("list", [])
    for i, c in enumerate(lst):
        if c["id"] == course_id:
            c.setdefault("session_details", []).append({
                "id": str(uuid.uuid4())[:8],
                "date": datetime.now().strftime("%Y-%m-%d"),
                "hours": 1,
                "topic": f"第{len(c.get('session_details', [])) + 1}次课",
                "attendance": 0,
            })
            _recalc_course_progress(c)
            lst[i] = c
            courses["list"] = lst
            _write_json(COURSES_FILE, courses)
            return APIResponse(success=True, data=c)
    raise HTTPException(status_code=404, detail="课程不存在")


# ═══════════════════════════════════════════════════════
# 课时明细管理（单节课 CRUD / 批量录入）
# ═══════════════════════════════════════════════════════

class SessionCreate(BaseModel):
    date: str = ""
    hours: int = 1
    topic: str = ""
    attendance: int = 0


class BatchSessions(BaseModel):
    items: list = []


def _recalc_course_progress(course: dict) -> dict:
    """基于 session_details 重算进度；无明细时用存储的 sessions 字段兜底。"""
    details = course.get("session_details", [])
    if details:
        total_sessions = sum(d.get("hours", 0) for d in details)
    else:
        total_sessions = course.get("sessions", 0)
    max_h = max(course.get("max_hours", 48), 1)
    progress = min(round(total_sessions / max_h * 100), 100)
    course["sessions"] = total_sessions
    course["progress"] = progress
    course["status"] = "已结课" if progress >= 100 else ("已过半" if progress >= 50 else "进行中")
    course["updated_at"] = datetime.now().isoformat()[:19]
    return course


@router.get("/courses/{course_id}/sessions", response_model=APIResponse)
async def list_sessions(course_id: str):
    courses = _read_json(COURSES_FILE)
    for c in courses.get("list", []):
        if c["id"] == course_id:
            details = sorted(c.get("session_details", []), key=lambda x: x.get("date", ""))
            return APIResponse(success=True, data={"items": details, "total": len(details)})
    raise HTTPException(status_code=404, detail="课程不存在")


@router.post("/courses/{course_id}/sessions", response_model=APIResponse)
async def create_session(course_id: str, data: SessionCreate):
    courses = _read_json(COURSES_FILE)
    for i, c in enumerate(courses.get("list", [])):
        if c["id"] == course_id:
            c.setdefault("session_details", []).append({
                "id": str(uuid.uuid4())[:8],
                "date": data.date or datetime.now().strftime("%Y-%m-%d"),
                "hours": max(1, data.hours),
                "topic": data.topic,
                "attendance": max(0, data.attendance),
            })
            _recalc_course_progress(c)
            courses["list"][i] = c
            _write_json(COURSES_FILE, courses)
            return APIResponse(success=True, message="已添加课时", data=c)
    raise HTTPException(status_code=404, detail="课程不存在")


@router.put("/courses/{course_id}/sessions/{session_id}", response_model=APIResponse)
async def update_session(course_id: str, session_id: str, data: SessionCreate):
    courses = _read_json(COURSES_FILE)
    for i, c in enumerate(courses.get("list", [])):
        if c["id"] == course_id:
            for d in c.get("session_details", []):
                if d["id"] == session_id:
                    d["date"] = data.date or d["date"]
                    d["hours"] = max(1, data.hours)
                    d["topic"] = data.topic
                    d["attendance"] = max(0, data.attendance)
                    break
            _recalc_course_progress(c)
            courses["list"][i] = c
            _write_json(COURSES_FILE, courses)
            return APIResponse(success=True, message="已更新课时", data=c)
    raise HTTPException(status_code=404, detail="课程不存在")


@router.delete("/courses/{course_id}/sessions/{session_id}", response_model=APIResponse)
async def delete_session(course_id: str, session_id: str):
    courses = _read_json(COURSES_FILE)
    for i, c in enumerate(courses.get("list", [])):
        if c["id"] == course_id:
            c["session_details"] = [d for d in c.get("session_details", []) if d["id"] != session_id]
            _recalc_course_progress(c)
            courses["list"][i] = c
            _write_json(COURSES_FILE, courses)
            return APIResponse(success=True, message="已删除课时", data=c)
    raise HTTPException(status_code=404, detail="课程不存在")


@router.post("/courses/{course_id}/sessions/batch", response_model=APIResponse)
async def batch_create_sessions(course_id: str, data: BatchSessions):
    courses = _read_json(COURSES_FILE)
    for i, c in enumerate(courses.get("list", [])):
        if c["id"] == course_id:
            for item in data.items:
                c.setdefault("session_details", []).append({
                    "id": str(uuid.uuid4())[:8],
                    "date": item.get("date", datetime.now().strftime("%Y-%m-%d")),
                    "hours": max(1, item.get("hours", 1)),
                    "topic": item.get("topic", ""),
                    "attendance": max(0, item.get("attendance", 0)),
                })
            _recalc_course_progress(c)
            courses["list"][i] = c
            _write_json(COURSES_FILE, courses)
            return APIResponse(success=True, message=f"已批量添加 {len(data.items)} 节课时", data=c)
    raise HTTPException(status_code=404, detail="课程不存在")


# ═══════════════════════════════════════════════════════
# 课程列表导出（Word）
# ═══════════════════════════════════════════════════════


class ExportRequest(BaseModel):
    course_ids: list[str] = []


def _build_course_export_docx(courses: list[dict]) -> bytes:
    """使用 python-docx 生成课程列表 Word 文档（含表格）。"""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # 标题
    title = doc.add_heading("课程列表导出", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 导出时间
    ts = doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}　　共 {len(courses)} 门课程")
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 空行

    # 表格
    headers = ["序号", "课程名称", "课程编号", "授课教师", "学期", "分类", "总课时", "已授课时", "进度", "学生数", "状态"]
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # 数据行
    for idx, c in enumerate(courses, 1):
        row = table.add_row()
        values = [
            str(idx),
            c.get("name", ""),
            c.get("code", ""),
            c.get("teacher", ""),
            c.get("semester", ""),
            c.get("category", ""),
            str(c.get("max_hours", "")),
            str(c.get("sessions", 0)),
            f"{c.get('progress', 0)}%",
            str(c.get("students", len(c.get("student_list", [])))),
            c.get("status", ""),
        ]
        for i, v in enumerate(values):
            row.cells[i].text = v
            for p in row.cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    # 列宽调整
    widths = [Cm(1.2), Cm(3.5), Cm(2.2), Cm(2.0), Cm(2.0), Cm(2.0), Cm(1.5), Cm(1.5), Cm(1.2), Cm(1.2), Cm(1.5)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/export")
async def export_courses(data: ExportRequest):
    """导出选中的课程列表为 Word (.docx) 文件。"""
    courses = _read_json(COURSES_FILE)
    all_courses = courses.get("list", [])

    if data.course_ids:
        selected = [c for c in all_courses if c["id"] in data.course_ids]
    else:
        selected = all_courses

    if not selected:
        raise HTTPException(status_code=400, detail="没有可导出的课程")

    file_data = _build_course_export_docx(selected)
    filename = f"课程列表_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    encoded_filename = quote(filename, safe="")

    return Response(
        content=file_data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
            "Content-Length": str(len(file_data)),
        },
    )


# ═══════════════════════════════════════════════════════
# 课程学生管理
# ═══════════════════════════════════════════════════════

class StudentCreate(BaseModel):
    name: str
    student_id: str
    class_name: str = ""


@router.get("/courses/{course_id}/students", response_model=APIResponse)
async def list_course_students(course_id: str, search: str = ""):
    """获取某门课的选课学生列表（含个人进度/出勤）。"""
    courses = _read_json(COURSES_FILE)
    for c in courses.get("list", []):
        if c["id"] == course_id:
            slist = c.get("student_list", [])
            if search:
                slist = [s for s in slist if search in s.get("name", "") or search in s.get("student_id", "")]
            return APIResponse(success=True, data={
                "items": slist,
                "total": len(slist),
                "studentCount": len(c.get("student_list", [])),
            })
    raise HTTPException(status_code=404, detail="课程不存在")


@router.post("/courses/{course_id}/students", response_model=APIResponse)
async def add_student(course_id: str, data: StudentCreate):
    """为课程添加学生（学号重复校验）。"""
    courses = _read_json(COURSES_FILE)
    for i, c in enumerate(courses.get("list", [])):
        if c["id"] == course_id:
            slist = c.get("student_list", [])
            # 学号重复校验
            if any(s.get("student_id") == data.student_id.strip() for s in slist):
                raise HTTPException(status_code=409, detail=f"学号「{data.student_id}」已存在")
            slist.append({
                "id": str(uuid.uuid4())[:8],
                "name": data.name.strip(),
                "student_id": data.student_id.strip(),
                "class": data.class_name.strip(),
                "progress": 0,
                "attendance": 0,
            })
            courses["list"][i]["student_list"] = slist
            courses["list"][i]["students"] = len(slist)
            courses["list"][i]["updated_at"] = datetime.now().isoformat()[:19]
            _write_json(COURSES_FILE, courses)
            return APIResponse(success=True, message=f"已添加学生「{data.name}」", data=courses["list"][i])
    raise HTTPException(status_code=404, detail="课程不存在")


@router.delete("/courses/{course_id}/students/{student_id}", response_model=APIResponse)
async def remove_student(course_id: str, student_id: str):
    """从课程移除学生。"""
    courses = _read_json(COURSES_FILE)
    for i, c in enumerate(courses.get("list", [])):
        if c["id"] == course_id:
            before = len(c.get("student_list", []))
            courses["list"][i]["student_list"] = [s for s in c.get("student_list", []) if s["id"] != student_id]
            if len(courses["list"][i]["student_list"]) == before:
                raise HTTPException(status_code=404, detail="学生不存在")
            courses["list"][i]["students"] = len(courses["list"][i]["student_list"])
            courses["list"][i]["updated_at"] = datetime.now().isoformat()[:19]
            _write_json(COURSES_FILE, courses)
            return APIResponse(success=True, message="已移除学生", data=courses["list"][i])
    raise HTTPException(status_code=404, detail="课程不存在")


# ── 班级编号 ────────────────────────────────────────

@router.post("/courses/{course_id}/generate-code", response_model=APIResponse)
async def generate_class_code(course_id: str):
    """生成/重置班级编号（10位字母数字）。"""
    import random, string
    courses = _read_json(COURSES_FILE)
    for i, c in enumerate(courses.get("list", [])):
        if c["id"] == course_id:
            code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=10))
            courses["list"][i]["class_code"] = code
            courses["list"][i]["updated_at"] = datetime.now().isoformat()[:19]
            _write_json(COURSES_FILE, courses)
            return APIResponse(success=True, message=f"班级编号已生成", data={"class_code": code})
    raise HTTPException(status_code=404, detail="课程不存在")


@router.post("/join-by-code", response_model=APIResponse)
async def join_by_code(data: dict):
    """学生通过班级编号加入课程。"""
    class_code = data.get("code", "").strip().upper()
    if not class_code:
        raise HTTPException(status_code=400, detail="请输入班级编号")

    student_name = data.get("student_name", "").strip()
    if not student_name:
        raise HTTPException(status_code=400, detail="请输入姓名")

    courses = _read_json(COURSES_FILE)
    for i, c in enumerate(courses.get("list", [])):
        if c.get("class_code", "").upper() == class_code:
            slist = c.get("student_list", [])
            if any(s.get("name") == student_name for s in slist):
                return APIResponse(success=True, message=f"你已在「{c['name']}」中", data=c)
            slist.append({
                "id": str(uuid.uuid4())[:8],
                "name": student_name,
                "student_id": "",
                "class": "",
                "progress": 0,
                "attendance": 0,
            })
            courses["list"][i]["student_list"] = slist
            courses["list"][i]["students"] = len(slist)
            courses["list"][i]["updated_at"] = datetime.now().isoformat()[:19]
            _write_json(COURSES_FILE, courses)
            return APIResponse(success=True, message=f"成功加入「{c['name']}」！", data=c)
    raise HTTPException(status_code=404, detail="班级编号无效，请检查后重试")