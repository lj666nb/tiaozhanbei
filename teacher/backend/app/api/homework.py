"""
作业批改与辅导 API — 智能批改、练习生成的 RESTful 接口（数据库持久化版）。
"""

from __future__ import annotations

import io
import json
import uuid
from datetime import datetime
from urllib.parse import quote

from fastapi import APIRouter, Depends, Form, HTTPException, Request, UploadFile
from fastapi.responses import Response
from sqlalchemy.orm import Session

from app.models.database import ExerciseBatch, HomeworkGrade, get_db
from app.models.schemas import (
    APIResponse,
    BatchGradingRequest,
    BatchGradingResponse,
    ExerciseRequest,
    ExerciseResponse,
    GradingResult,
    HomeworkSubmission,
)
from app.services.homework_service import generate_exercises as gen_exercises
from app.services.homework_service import grade_batch, grade_submission, process_uploaded_file

router = APIRouter(prefix="/api/homework", tags=["作业批改"])


@router.post("/grade", response_model=APIResponse)
async def grade_homework(submission: HomeworkSubmission, db: Session = Depends(get_db)):
    """批改单个作业并保存结果。"""
    try:
        result = grade_submission(submission)

        # 保存到数据库
        record = HomeworkGrade(
            id=str(uuid.uuid4())[:12],
            student_name=submission.student_name,
            course_name=submission.course_name,
            chapter=submission.chapter or "",
            question_text=submission.question_text,
            student_answer=submission.student_answer,
            question_type=submission.question_type or "主观题",
            max_score=submission.max_score or 100,
            score=result.score,
            percentage=result.percentage,
            feedback=result.feedback or "",
            strengths=json.dumps(result.strengths, ensure_ascii=False),
            weaknesses=json.dumps(result.weaknesses, ensure_ascii=False),
            suggestions=json.dumps(result.suggestions, ensure_ascii=False),
            knowledge_points=json.dumps(result.knowledge_points, ensure_ascii=False),
            detailed_analysis=result.detailed_analysis or "",
            created_at=datetime.now(),
        )
        db.add(record)
        db.commit()

        return APIResponse(
            success=True,
            message="批改完成",
            data=result.model_dump() | {"id": record.id},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/batch-grade", response_model=APIResponse)
async def batch_grade(request: BatchGradingRequest, db: Session = Depends(get_db)):
    """批量批改作业并保存全部结果。"""
    try:
        results, avg_score, distribution = grade_batch(request.submissions)

        # 生成批次ID
        batch_id = str(uuid.uuid4())[:12]
        # 从第一条提交提取来源信息
        first_sub = request.submissions[0] if request.submissions else None
        source_file = getattr(request, 'source_file', '') or ''

        # 批量保存
        now = datetime.now()
        for i, s in enumerate(request.submissions):
            r = results[i] if i < len(results) else None
            if r:
                record = HomeworkGrade(
                    id=str(uuid.uuid4())[:12],
                    student_name=s.student_name,
                    course_name=s.course_name,
                    chapter=s.chapter or "",
                    question_text=s.question_text,
                    student_answer=s.student_answer,
                    question_type=s.question_type or "主观题",
                    max_score=s.max_score or 100,
                    score=r.score,
                    percentage=r.percentage,
                    feedback=r.feedback or "",
                    strengths=json.dumps(r.strengths, ensure_ascii=False),
                    weaknesses=json.dumps(r.weaknesses, ensure_ascii=False),
                    suggestions=json.dumps(r.suggestions, ensure_ascii=False),
                    knowledge_points=json.dumps(r.knowledge_points, ensure_ascii=False),
                    detailed_analysis=r.detailed_analysis or "",
                    source_file=source_file,
                    batch_id=batch_id,
                    created_at=now,
                )
                db.add(record)
        db.commit()

        return APIResponse(
            success=True,
            message=f"共批改 {len(results)} 份作业，平均分 {avg_score}",
            data=BatchGradingResponse(
                results=results,
                total_submissions=len(results),
                avg_score=avg_score,
                class_distribution=distribution,
            ).model_dump() | {"batch_id": batch_id},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/exercises", response_model=APIResponse)
async def generate_exercises_api(request: ExerciseRequest, db: Session = Depends(get_db)):
    """生成练习题并保存。"""
    try:
        result = gen_exercises(request)

        # 保存到数据库
        batch = ExerciseBatch(
            id=str(uuid.uuid4())[:12],
            course_name=request.course_name,
            chapter=request.chapter or "",
            difficulty=request.difficulty or "中等",
            total=result.total,
            exercises_json=json.dumps([e.model_dump() for e in result.exercises], ensure_ascii=False),
            created_at=datetime.now(),
        )
        db.add(batch)
        db.commit()

        return APIResponse(
            success=True,
            message=f"已生成 {result.total} 道练习题",
            data=result.model_dump() | {"batch_id": batch.id},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── 文件上传批改 ──────────────────────────────────────

@router.post("/upload", response_model=APIResponse)
async def upload_homework_file(
    file: UploadFile,
    course: str = Form(""),
    parse_only: bool = Form(False),
):
    """上传作业文件进行批改。

    支持 PDF / Word / 图片（JPG/PNG/WebP）/ CSV / TXT。
    后端自动提取文本 → 解析结构 → AI 批改。
    parse_only=True 时只解析文本不批改，适用于有独立答案文件的场景。
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="请选择文件")

    # 检查文件大小（50MB）
    MAX_SIZE = 50 * 1024 * 1024
    content = await file.read()
    if len(content) > MAX_SIZE:
        raise HTTPException(status_code=413, detail="文件大小超过 50MB 限制")

    try:
        results, raw_text = process_uploaded_file(content, file.filename, course, parse_only=parse_only)
        return APIResponse(
            success=True,
            message=f"批改完成，共 {len(results)} 份作业",
            data={
                "filename": file.filename,
                "results": results,
                "raw_text": raw_text,
                "total": len(results),
            },
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")


@router.post("/batch-upload", response_model=APIResponse)
async def batch_upload_homework_files(
    files: list[UploadFile],
    course: str = Form(""),
):
    """批量上传多个作业文件进行批改。"""
    all_results = []
    for file in files:
        if not file.filename:
            continue
        content = await file.read()
        try:
            results, _ = process_uploaded_file(content, file.filename, course)
            for r in results:
                r["source_file"] = file.filename
            all_results.extend(results)
        except Exception as e:
            all_results.append({
                "student_name": "未知学生",
                "source_file": file.filename,
                "score": 0, "max_score": 100, "percentage": 0,
                "feedback": f"处理失败: {str(e)[:100]}",
                "strengths": [], "weaknesses": [], "suggestions": [],
                "knowledge_points": [], "detailed_analysis": "",
            })

    return APIResponse(
        success=len(all_results) > 0,
        message=f"共处理 {len(files)} 个文件，批改 {len(all_results)} 份作业",
        data={"results": all_results, "total": len(all_results)},
    )


# ── 练习题导出 Word ────────────────────────────────────


@router.post("/exercises/export-word")
async def export_exercises_word(request: Request):
    """将生成的练习题导出为 Word (.docx) 文档。"""
    body = await request.json()
    exercises = body.get("exercises", [])
    course_name = body.get("course_name", "练习题")
    chapter = body.get("chapter", "")

    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 页面标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(f"练习题 - {course_name}")
        run.bold = True
        run.font.size = Pt(22)

        if chapter:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = sub.add_run(f"章节：{chapter}")
            run2.font.size = Pt(14)
            run2.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph(f"共 {len(exercises)} 道练习题")
        doc.add_paragraph("")

        for i, ex in enumerate(exercises):
            q_type = ex.get("type", "")
            diff = ex.get("difficulty", "")
            heading = doc.add_heading(f"第{i+1}题 [{q_type}] ({diff})", level=2)

            # 题目
            doc.add_paragraph(ex.get("question", ""))

            # 选项
            options = ex.get("options", [])
            if options:
                for opt in options:
                    doc.add_paragraph(opt, style="List Bullet")

            # 元信息
            meta = doc.add_paragraph()
            meta_run = meta.add_run(f"知识点：{ex.get('knowledge_point', '')}　｜　预计用时：{ex.get('estimated_time', 5)} 分钟")
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = RGBColor(120, 120, 120)

            # 答案
            ans_heading = doc.add_heading("答案", level=3)
            doc.add_paragraph(ex.get("answer", ""))

            # 解析
            exp_heading = doc.add_heading("解析", level=3)
            doc.add_paragraph(ex.get("explanation", ""))

            doc.add_paragraph("")

        # 写入内存
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"练习题_{course_name}"
        if chapter:
            filename += f"_{chapter}"
        filename += ".docx"

        from urllib.parse import quote as _quote
        encoded = _quote(filename, safe="")
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded}",
            },
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装，无法生成 Word 文档")


# ── 批次管理 ──────────────────────────────────────

@router.get("/batches", response_model=APIResponse)
async def list_batches(course: str = "", db: Session = Depends(get_db)):
    """获取批次列表（按 batch_id + source_file 聚合）。"""
    from sqlalchemy import func
    query = db.query(
        HomeworkGrade.batch_id,
        HomeworkGrade.source_file,
        func.count(HomeworkGrade.id).label("total"),
        func.avg(HomeworkGrade.percentage).label("avg_percentage"),
        func.max(HomeworkGrade.created_at).label("latest_time"),
    ).filter(
        HomeworkGrade.id != "_SEED_DONE_",
        HomeworkGrade.batch_id != "",
    )
    if course:
        query = query.filter(HomeworkGrade.course_name == course)
    query = query.group_by(HomeworkGrade.batch_id, HomeworkGrade.source_file)\
        .order_by(func.max(HomeworkGrade.created_at).desc()).limit(50)
    rows = query.all()

    batches = []
    for r in rows:
        batches.append({
            "batch_id": r.batch_id,
            "source_file": r.source_file or "未命名文件",
            "total": r.total,
            "avg_percentage": round(float(r.avg_percentage), 1) if r.avg_percentage else 0,
            "latest_time": r.latest_time.isoformat() if r.latest_time else "",
        })
    return APIResponse(success=True, data={"batches": batches, "total": len(batches)})


@router.get("/batches/{batch_id}", response_model=APIResponse)
async def get_batch(batch_id: str, db: Session = Depends(get_db)):
    """获取某批次的所有记录详情。"""
    records = db.query(HomeworkGrade).filter(
        HomeworkGrade.batch_id == batch_id,
        HomeworkGrade.id != "_SEED_DONE_",
    ).order_by(HomeworkGrade.created_at.asc()).all()
    if not records:
        raise HTTPException(status_code=404, detail="批次不存在")
    return APIResponse(success=True, data={
        "batch_id": batch_id,
        "source_file": records[0].source_file or "未命名文件",
        "total": len(records),
        "records": [r.to_dict() for r in records],
    })


@router.delete("/batches/{batch_id}", response_model=APIResponse)
async def delete_batch(batch_id: str, db: Session = Depends(get_db)):
    """删除整个批次的所有记录。"""
    count = db.query(HomeworkGrade).filter(
        HomeworkGrade.batch_id == batch_id,
        HomeworkGrade.id != "_SEED_DONE_",
    ).delete(synchronize_session=False)
    db.commit()
    return APIResponse(success=True, message=f"已删除批次 {batch_id}（{count} 条记录）")


# ── 历史记录查询 ──────────────────────────────────────

@router.get("/grades", response_model=APIResponse)
async def list_grades(course: str = "", archived: bool = None, db: Session = Depends(get_db)):
    """获取批改历史记录。

    - archived 不传：返回全部记录（作业批改页面使用）
    - archived=true：只返回已手动归档的记录（教学台账使用）
    """
    query = db.query(HomeworkGrade).filter(
        HomeworkGrade.id != "_SEED_DONE_",
        HomeworkGrade.question_type != "手动录入",
    )
    if archived is not None:
        query = query.filter(HomeworkGrade.is_archived == archived)
    if course:
        query = query.filter(HomeworkGrade.course_name == course)
    records = query.order_by(HomeworkGrade.created_at.desc()).limit(100).all()
    return APIResponse(success=True, data={
        "total": len(records),
        "items": [r.to_dict() for r in records],
    })


@router.post("/archive", response_model=APIResponse)
async def archive_grades(request: Request, db: Session = Depends(get_db)):
    """将批改结果归档至教学台账（直接保存，不重复AI批改）。"""
    body = await request.json()
    results = body.get("results", [])
    if not results:
        raise HTTPException(status_code=400, detail="没有可归档的结果")
    saved = 0
    skipped = 0
    for r in results:
        student_name = r.get("student_name", "").strip() or "未知学生"
        course_name = r.get("course_name", "").strip()
        question_text = r.get("question_text", "").strip()
        student_answer = r.get("student_answer", "").strip()

        # 跳过无效记录（没有学生姓名或课程名称的）
        if not course_name and not question_text:
            skipped += 1
            continue

        # 安全处理列表/JSON字段
        def _safe_json(val, default="[]"):
            if val is None:
                return default
            if isinstance(val, str):
                return val
            return json.dumps(val, ensure_ascii=False)

        record = HomeworkGrade(
            id=str(uuid.uuid4())[:12],
            student_name=student_name,
            course_name=course_name,
            chapter=r.get("chapter", ""),
            question_text=question_text,
            student_answer=student_answer,
            question_type=r.get("question_type", "主观题"),
            max_score=float(r.get("max_score", 100)),
            score=float(r.get("score", 0)),
            percentage=float(r.get("percentage", 0)),
            feedback=r.get("feedback", ""),
            strengths=_safe_json(r.get("strengths")),
            weaknesses=_safe_json(r.get("weaknesses")),
            suggestions=_safe_json(r.get("suggestions")),
            knowledge_points=_safe_json(r.get("knowledge_points")),
            detailed_analysis=r.get("detailed_analysis", ""),
            source_file=r.get("source_file", r.get("_sourceFile", "")),
            batch_id=r.get("batch_id", ""),
            is_archived=True,
        )
        db.add(record)
        saved += 1
    db.commit()
    msg = f"已归档 {saved} 条批改结果至教学台账"
    if skipped > 0:
        msg += f"（跳过 {skipped} 条空记录）"
    return APIResponse(success=True, message=msg)


@router.get("/exercises/list", response_model=APIResponse)
async def list_exercises(course: str = "", db: Session = Depends(get_db)):
    """获取出题历史。"""
    query = db.query(ExerciseBatch)
    if course:
        query = query.filter(ExerciseBatch.course_name == course)
    batches = query.order_by(ExerciseBatch.created_at.desc()).limit(50).all()
    return APIResponse(success=True, data={
        "total": len(batches),
        "items": [b.to_dict() for b in batches],
    })


@router.get("/grades/{grade_id}", response_model=APIResponse)
async def get_grade(grade_id: str, db: Session = Depends(get_db)):
    """获取单条批改详情。"""
    record = db.query(HomeworkGrade).filter(HomeworkGrade.id == grade_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="记录不存在")
    return APIResponse(success=True, data=record.to_dict())


@router.delete("/grades/{grade_id}", response_model=APIResponse)
async def delete_grade(grade_id: str, db: Session = Depends(get_db)):
    """删除单条批改记录。"""
    record = db.query(HomeworkGrade).filter(HomeworkGrade.id == grade_id).first()
    if not record:
        raise HTTPException(status_code=404, detail="记录不存在")
    db.delete(record)
    db.commit()
    return APIResponse(success=True, message="批改记录已删除")
